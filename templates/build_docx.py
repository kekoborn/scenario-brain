#!/usr/bin/env python3
"""
Сборка финального сценария в .docx из scenario-final.md + выбранного названия.
Использование:
    python3 build_docx.py <scenario-final.md> "<Название видео>" <output.docx> [формат]

формат: suffler (по умолчанию) — текст ведущего слово-в-слово для телепромптера;
        oporny — опорные тезисы (первая фраза каждого абзаца буллетом).

Парсит блоки (## БЛОК ...) и раздел КЛЮЧЕВЫЕ ФРАЗЫ. Карту/источники в docx не включает.
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    md_path, title, out_path = sys.argv[1], sys.argv[2], sys.argv[3]
    fmt = sys.argv[4] if len(sys.argv) > 4 else "suffler"  # suffler | oporny
    lines = open(md_path, encoding="utf-8").read().splitlines()

    doc = Document()
    # базовый шрифт
    style = doc.styles["Normal"].font
    style.name = "Georgia"; style.size = Pt(12)

    # титул
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title); r.bold = True; r.font.size = Pt(22)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Сценарий видео · обучающий формат · стиль М. Дашкиева")
    rs.italic = True; rs.font.size = Pt(11); rs.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_paragraph()

    def strip_md(s):
        s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
        s = s.replace("> ", "").replace("«","«").strip()
        return s

    # найти старт первого блока
    start = next((i for i,l in enumerate(lines) if l.startswith("## БЛОК")), 0)
    buf = []
    in_keys = False

    def flush():
        nonlocal buf
        if not buf:
            return
        para = " ".join(buf).strip()
        if fmt == "oporny":
            # опорный: первая фраза абзаца как тезис-буллет
            thesis = re.split(r"(?<=[.!?])\s", para)[0].strip()
            try:
                doc.add_paragraph(thesis, style="List Bullet")
            except KeyError:
                doc.add_paragraph("• " + thesis)
        else:
            doc.add_paragraph(para)  # суфлёр: полный текст слово-в-слово
        buf = []

    for l in lines[start:]:
        s = l.rstrip()
        if s.startswith("# КАРТА") or s.startswith("# ИСТОЧ"):
            flush(); break
        if s.startswith("## БЛОК"):
            flush()
            h = doc.add_heading(strip_md(s[3:]), level=1)
            continue
        if s.startswith("# КЛЮЧЕВЫЕ"):
            flush()
            doc.add_page_break()
            doc.add_heading("Ключевые фразы", level=1)
            in_keys = True
            continue
        if s.strip() == "---" or s.strip() == "":
            flush(); continue
        if in_keys and re.match(r"^\d+\.", s):
            flush()
            doc.add_paragraph(strip_md(s), style="List Number" if False else None)
            p = doc.paragraphs[-1]
            continue
        buf.append(s.strip())
    flush()

    doc.save(out_path)
    print("saved:", out_path)

if __name__ == "__main__":
    main()
